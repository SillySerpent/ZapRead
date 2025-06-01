import os
import re
from docx import Document
from .utils import process_text_bionic, create_temp_output_path, debug_print

def create_bionic_docx_file(input_path):
    """
    Create a bionic reading DOCX file by modifying the original document.
    
    Args:
        input_path (str): Path to the input DOCX file.
        
    Returns:
        str: Path to the output DOCX file.
    """
    try:
        # Open the document
        doc = Document(input_path)
        
        # Create output path
        temp_dir, output_path = create_temp_output_path(input_path, 'docx')
        
        # Process each paragraph
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                _process_paragraph_runs(paragraph)
        
        # Process tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            _process_paragraph_runs(paragraph)
        
        # Save the modified document
        doc.save(output_path)
        
        debug_print(f"DOCX file processed successfully: {output_path}")
        return output_path
        
    except Exception as e:
        print(f"Error processing DOCX file: {e}")
        raise

def _process_paragraph_runs(paragraph):
    """
    Process runs within a paragraph to apply bionic reading formatting.
    Enhanced to handle mathematical content properly.
    
    Args:
        paragraph: Document paragraph object to process.
    """
    # Import here to avoid circular imports
    from .utils.pattern_matcher import split_preserving_math
    from .utils.math_detector import should_preserve_as_math
    
    # Create a list to store new runs
    new_runs_data = []
    
    # Process each run
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
            
        # Store run formatting
        font = run.font
        
        # Use enhanced text segmentation
        segments = split_preserving_math(text)
        
        for segment, is_word, is_math in segments:
            if is_word and not is_math:
                # Process word for bionic reading
                clean_word, bold_part, regular_part = process_text_bionic(segment)
                
                if bold_part:
                    # Add bold part as separate run
                    new_runs_data.append({
                        'text': bold_part,
                        'bold': True,  # Make it bold
                        'italic': font.italic,
                        'size': font.size,
                        'color': font.color.rgb if font.color and font.color.rgb else None,
                        'name': font.name
                    })
                
                if regular_part:
                    # Add regular part
                    new_runs_data.append({
                        'text': regular_part,
                        'bold': font.bold,  # Keep original bold setting
                        'italic': font.italic,
                        'size': font.size,
                        'color': font.color.rgb if font.color and font.color.rgb else None,
                        'name': font.name
                    })
            else:
                # Preserve math, spaces, and punctuation as-is
                new_runs_data.append({
                    'text': segment,
                    'bold': font.bold,
                    'italic': font.italic,
                    'size': font.size,
                    'color': font.color.rgb if font.color and font.color.rgb else None,
                    'name': font.name
                })
    
    # Clear existing runs
    paragraph.clear()
    
    # Add new runs with bionic formatting
    for run_data in new_runs_data:
        new_run = paragraph.add_run(run_data['text'])
        
        # Apply formatting
        if run_data['bold'] is not None:
            new_run.font.bold = run_data['bold']
        if run_data['italic'] is not None:
            new_run.font.italic = run_data['italic']
        if run_data['size'] is not None:
            new_run.font.size = run_data['size']
        if run_data['color'] is not None:
            new_run.font.color.rgb = run_data['color']
        if run_data['name'] is not None:
            new_run.font.name = run_data['name']

def process_docx_file(input_path):
    """
    Main entry point for DOCX file processing.
    
    Args:
        input_path (str): Path to the input DOCX file.
        
    Returns:
        dict: Processing result with success status and output path.
    """
    try:
        output_path = create_bionic_docx_file(input_path)
        return {
            'success': True,
            'output_path': output_path,
            'file_type': 'docx'
        }
    except Exception as e:
        return {
            'success': False,
            'error': str(e),
            'file_type': 'docx'
        } 