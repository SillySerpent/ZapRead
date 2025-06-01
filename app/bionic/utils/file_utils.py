"""
File Utilities for Bionic Processor

Provides comprehensive file handling capabilities including type detection,
validation, metadata extraction, and path management.
"""

import os
import mimetypes
import magic
from pathlib import Path
from typing import Dict, Any, Optional, Union, List
import logging

logger = logging.getLogger(__name__)


def detect_file_type(file_path: str) -> str:
    """
    Detect file type with multiple methods for accuracy.
    
    Args:
        file_path: Path to the file
        
    Returns:
        Detected file type (lowercase extension without dot)
    """
    try:
        # Method 1: File extension
        ext_type = Path(file_path).suffix.lower().lstrip('.')
        
        # Method 2: MIME type
        mime_type, _ = mimetypes.guess_type(file_path)
        
        # Method 3: Magic number detection (most reliable)
        try:
            file_magic = magic.from_file(file_path, mime=True)
        except Exception as e:
            logger.debug(f"Magic detection failed for {file_path}: {e}")
            file_magic = None
        
        # Mapping of MIME types to file types
        mime_mappings = {
            'application/pdf': 'pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
            'application/msword': 'doc',
            'text/plain': 'txt',
            'text/html': 'html',
            'text/markdown': 'md',
            'application/rtf': 'rtf'
        }
        
        # Priority: Magic > MIME > Extension
        if file_magic and file_magic in mime_mappings:
            detected_type = mime_mappings[file_magic]
            logger.debug(f"Magic detection: {file_path} -> {detected_type}")
            return detected_type
        
        if mime_type and mime_type in mime_mappings:
            detected_type = mime_mappings[mime_type]
            logger.debug(f"MIME detection: {file_path} -> {detected_type}")
            return detected_type
        
        if ext_type:
            logger.debug(f"Extension detection: {file_path} -> {ext_type}")
            return ext_type
        
        # Fallback for specific cases
        if file_magic:
            if 'pdf' in file_magic.lower():
                return 'pdf'
            elif 'word' in file_magic.lower() or 'openxml' in file_magic.lower():
                return 'docx'
            elif 'text' in file_magic.lower():
                return 'txt'
        
        logger.warning(f"Could not detect file type for {file_path}")
        return 'unknown'
        
    except Exception as e:
        logger.error(f"File type detection failed for {file_path}: {e}")
        return 'unknown'


def validate_file(file_path: str, config: Any) -> Dict[str, Any]:
    """
    Validate file and extract metadata.
    
    Args:
        file_path: Path to the file
        config: Configuration object with validation settings
        
    Returns:
        Dictionary with file information and validation results
    """
    file_info = {
        'file_path': file_path,
        'exists': False,
        'size_bytes': 0,
        'size_mb': 0.0,
        'file_type': 'unknown',
        'is_valid': False,
        'validation_errors': [],
        'metadata': {}
    }
    
    try:
        # Check file existence
        if not os.path.exists(file_path):
            file_info['validation_errors'].append(f"File does not exist: {file_path}")
            return file_info
        
        file_info['exists'] = True
        
        # Get file size
        size_bytes = os.path.getsize(file_path)
        file_info['size_bytes'] = size_bytes
        file_info['size_mb'] = size_bytes / (1024 * 1024)
        
        # Check file size limits
        max_size_mb = getattr(config, 'max_file_size_mb', 500)
        if file_info['size_mb'] > max_size_mb:
            file_info['validation_errors'].append(
                f"File too large: {file_info['size_mb']:.1f}MB > {max_size_mb}MB"
            )
        
        # Detect file type
        file_type = detect_file_type(file_path)
        file_info['file_type'] = file_type
        
        # Check if file type is supported
        supported_types = getattr(config, 'supported_types', ['pdf', 'docx', 'txt'])
        if file_type not in supported_types:
            file_info['validation_errors'].append(
                f"Unsupported file type: {file_type}. Supported: {supported_types}"
            )
        
        # Extract file-specific metadata
        file_info['metadata'] = extract_file_metadata(file_path, file_type)
        
        # File is valid if no errors
        file_info['is_valid'] = len(file_info['validation_errors']) == 0
        
    except Exception as e:
        file_info['validation_errors'].append(f"Validation error: {str(e)}")
        logger.error(f"File validation failed for {file_path}: {e}")
    
    return file_info


def extract_file_metadata(file_path: str, file_type: str) -> Dict[str, Any]:
    """
    Extract metadata specific to file type.
    
    Args:
        file_path: Path to the file
        file_type: Detected file type
        
    Returns:
        Dictionary with file-specific metadata
    """
    metadata = {
        'created_time': None,
        'modified_time': None,
        'page_count': 0,
        'line_count': 0,
        'word_count': 0,
        'character_count': 0
    }
    
    try:
        # Basic file stats
        stat = os.stat(file_path)
        metadata['created_time'] = stat.st_ctime
        metadata['modified_time'] = stat.st_mtime
        
        if file_type == 'pdf':
            metadata.update(extract_pdf_metadata(file_path))
        elif file_type == 'docx':
            metadata.update(extract_docx_metadata(file_path))
        elif file_type == 'txt':
            metadata.update(extract_text_metadata(file_path))
    
    except Exception as e:
        logger.warning(f"Metadata extraction failed for {file_path}: {e}")
    
    return metadata


def extract_pdf_metadata(file_path: str) -> Dict[str, Any]:
    """Extract PDF-specific metadata."""
    metadata = {}
    
    try:
        import fitz  # PyMuPDF
        
        with fitz.open(file_path) as doc:
            metadata.update({
                'page_count': len(doc),
                'title': doc.metadata.get('title', ''),
                'author': doc.metadata.get('author', ''),
                'subject': doc.metadata.get('subject', ''),
                'creator': doc.metadata.get('creator', ''),
                'producer': doc.metadata.get('producer', ''),
                'creation_date': doc.metadata.get('creationDate', ''),
                'modification_date': doc.metadata.get('modDate', ''),
                'encrypted': doc.needs_pass,
                'has_forms': doc.is_form_pdf,
                'has_annotations': any(page.annots() for page in doc),
                'has_images': False,  # Will be updated during text extraction
                'total_text_length': 0
            })
            
            # Count total text and check for images
            total_text = 0
            has_images = False
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                total_text += len(text)
                
                if not has_images and page.get_images():
                    has_images = True
            
            metadata['total_text_length'] = total_text
            metadata['has_images'] = has_images
            metadata['average_text_per_page'] = total_text / len(doc) if len(doc) > 0 else 0
    
    except ImportError:
        logger.warning("PyMuPDF not available for PDF metadata extraction")
    except Exception as e:
        logger.warning(f"PDF metadata extraction failed: {e}")
    
    return metadata


def extract_docx_metadata(file_path: str) -> Dict[str, Any]:
    """Extract DOCX-specific metadata."""
    metadata = {}
    
    try:
        from docx import Document
        
        doc = Document(file_path)
        
        # Basic document info
        metadata.update({
            'paragraph_count': len(doc.paragraphs),
            'table_count': len(doc.tables),
            'section_count': len(doc.sections),
            'has_headers': any(section.header for section in doc.sections),
            'has_footers': any(section.footer for section in doc.sections),
        })
        
        # Count text content
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        full_text = '\n'.join(text_content)
        metadata.update({
            'word_count': len(full_text.split()),
            'character_count': len(full_text),
            'line_count': len(text_content)
        })
        
        # Document properties (if available)
        if hasattr(doc.core_properties, 'title'):
            metadata['title'] = doc.core_properties.title or ''
        if hasattr(doc.core_properties, 'author'):
            metadata['author'] = doc.core_properties.author or ''
        if hasattr(doc.core_properties, 'subject'):
            metadata['subject'] = doc.core_properties.subject or ''
    
    except ImportError:
        logger.warning("python-docx not available for DOCX metadata extraction")
    except Exception as e:
        logger.warning(f"DOCX metadata extraction failed: {e}")
    
    return metadata


def extract_text_metadata(file_path: str) -> Dict[str, Any]:
    """Extract text file metadata."""
    metadata = {}
    
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        lines = content.split('\n')
        words = content.split()
        
        metadata.update({
            'line_count': len(lines),
            'word_count': len(words),
            'character_count': len(content),
            'non_empty_lines': sum(1 for line in lines if line.strip()),
            'encoding': 'utf-8',  # Assumed since we opened with utf-8
            'has_unicode': any(ord(char) > 127 for char in content)
        })
    
    except Exception as e:
        logger.warning(f"Text metadata extraction failed: {e}")
    
    return metadata


def create_output_path(input_path: str, output_dir: Optional[str] = None, 
                      suffix: str = '_bionic', extension: str = 'html') -> str:
    """
    Create standardized output path for processed files.
    
    Args:
        input_path: Original file path
        output_dir: Output directory (optional)
        suffix: Suffix to add to filename
        extension: Output file extension
        
    Returns:
        Generated output path
    """
    input_path_obj = Path(input_path)
    
    # Create output filename
    output_filename = f"{input_path_obj.stem}{suffix}.{extension}"
    
    if output_dir:
        output_path = Path(output_dir) / output_filename
        # Create output directory if it doesn't exist
        os.makedirs(output_dir, exist_ok=True)
    else:
        output_path = input_path_obj.parent / output_filename
    
    return str(output_path)


def ensure_directory(directory_path: str) -> bool:
    """
    Ensure directory exists, create if necessary.
    
    Args:
        directory_path: Path to directory
        
    Returns:
        True if directory exists/was created, False otherwise
    """
    try:
        os.makedirs(directory_path, exist_ok=True)
        return True
    except Exception as e:
        logger.error(f"Failed to create directory {directory_path}: {e}")
        return False


def clean_filename(filename: str) -> str:
    """
    Clean filename by removing/replacing invalid characters.
    
    Args:
        filename: Original filename
        
    Returns:
        Cleaned filename safe for filesystem
    """
    import re
    
    # Remove/replace invalid characters
    cleaned = re.sub(r'[<>:"/\\|?*]', '_', filename)
    
    # Remove multiple consecutive underscores
    cleaned = re.sub(r'_+', '_', cleaned)
    
    # Remove leading/trailing underscores and dots
    cleaned = cleaned.strip('_.')
    
    # Ensure filename is not empty
    if not cleaned:
        cleaned = 'processed_file'
    
    return cleaned


def get_temp_file_path(base_name: str = 'bionic_temp', extension: str = 'tmp',
                       temp_dir: Optional[str] = None) -> str:
    """
    Generate temporary file path.
    
    Args:
        base_name: Base name for temporary file
        extension: File extension
        temp_dir: Temporary directory (uses system temp if None)
        
    Returns:
        Path to temporary file
    """
    import tempfile
    import uuid
    
    if temp_dir is None:
        temp_dir = tempfile.gettempdir()
    
    # Generate unique filename
    unique_id = str(uuid.uuid4())[:8]
    temp_filename = f"{base_name}_{unique_id}.{extension}"
    
    return os.path.join(temp_dir, temp_filename)


def copy_file_with_validation(source: str, destination: str) -> bool:
    """
    Copy file with validation.
    
    Args:
        source: Source file path
        destination: Destination file path
        
    Returns:
        True if copy successful, False otherwise
    """
    try:
        import shutil
        
        # Validate source exists
        if not os.path.exists(source):
            logger.error(f"Source file does not exist: {source}")
            return False
        
        # Create destination directory if needed
        dest_dir = os.path.dirname(destination)
        if dest_dir and not ensure_directory(dest_dir):
            return False
        
        # Copy file
        shutil.copy2(source, destination)
        
        # Validate copy
        if os.path.exists(destination):
            source_size = os.path.getsize(source)
            dest_size = os.path.getsize(destination)
            
            if source_size == dest_size:
                logger.debug(f"Successfully copied {source} to {destination}")
                return True
            else:
                logger.error(f"File copy size mismatch: {source_size} != {dest_size}")
                return False
        else:
            logger.error(f"Destination file not created: {destination}")
            return False
    
    except Exception as e:
        logger.error(f"File copy failed: {e}")
        return False


def cleanup_temp_files(file_patterns: List[str], temp_dir: Optional[str] = None) -> int:
    """
    Clean up temporary files matching patterns.
    
    Args:
        file_patterns: List of file patterns to match
        temp_dir: Directory to clean (uses system temp if None)
        
    Returns:
        Number of files cleaned up
    """
    import glob
    import tempfile
    
    if temp_dir is None:
        temp_dir = tempfile.gettempdir()
    
    cleaned_count = 0
    
    for pattern in file_patterns:
        full_pattern = os.path.join(temp_dir, pattern)
        matching_files = glob.glob(full_pattern)
        
        for file_path in matching_files:
            try:
                os.remove(file_path)
                cleaned_count += 1
                logger.debug(f"Cleaned up temporary file: {file_path}")
            except Exception as e:
                logger.warning(f"Failed to clean up {file_path}: {e}")
    
    return cleaned_count 