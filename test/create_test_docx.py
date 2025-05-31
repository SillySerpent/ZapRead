from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_test_docx(output_path):
    doc = Document()
    
    # Add title
    title = doc.add_paragraph()
    title_run = title.add_run("Test DOCX Document with Formatting")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add introduction
    doc.add_paragraph("This is a test DOCX document created to demonstrate various formatting features including headings, paragraphs, and text formatting.")
    
    # Add heading
    heading1 = doc.add_heading("Section 1: Different Text Formats", level=2)
    
    # Regular paragraph
    doc.add_paragraph("This is a regular paragraph with standard text formatting. It should be preserved in the bionic reading version.")
    
    # Bold and italic text
    p = doc.add_paragraph()
    p.add_run("This contains ")
    p.add_run("bold text").bold = True
    p.add_run(" and ")
    p.add_run("italic text").italic = True
    p.add_run(" and ")
    r = p.add_run("both bold and italic")
    r.bold = True
    r.italic = True
    p.add_run(" which should be preserved.")
    
    # Another heading
    heading2 = doc.add_heading("Section 2: Structure and Spacing", level=2)
    
    # Multiple paragraphs
    doc.add_paragraph("This is the first paragraph in a section. It demonstrates paragraph spacing and structure preservation.")
    doc.add_paragraph("This is a second paragraph with different content. The spacing between paragraphs should be maintained in the bionic reading version.")
    
    p = doc.add_paragraph("This third paragraph has some ")
    p.add_run("important words").bold = True
    p.add_run(" that should be highlighted. The bionic reading algorithm should still work correctly with this formatting.")
    
    # Conclusion
    doc.add_heading("Conclusion", level=2)
    doc.add_paragraph("This test DOCX demonstrates various formatting features that should be preserved during bionic reading conversion. The structure, headings, paragraph spacing, and text formatting should all be maintained.")
    
    # Save the document
    doc.save(output_path)
    
    return output_path

if __name__ == "__main__":
    output_file = "test_formatted.docx"
    create_test_docx(output_file)
    print(f"Test DOCX created successfully: {output_file}") 