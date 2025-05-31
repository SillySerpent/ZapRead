import os
import re
import fitz  # pymupdf
import tempfile
import magic
import base64
from docx import Document
from docx.shared import RGBColor
from docx.oxml.shared import qn
import pdfplumber
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.colors import black
import io
from PIL import Image
import json
from reportlab.lib.utils import ImageReader

# Debug settings
DEBUG = False  # Set to False for production use

def debug_print(*args, **kwargs):
    """
    Print debug information only if DEBUG is enabled.
    """
    if DEBUG:
        print(*args, **kwargs)

def process_text_bionic(text):
    """
    Process text to create bionic reading format with bold prefixes.
    
    Args:
        text (str): The text to process.
        
    Returns:
        tuple: (original_text, bold_part, regular_part) for reconstruction
    """
    if not text or not text.strip():
        return text, "", ""
    
    # Clean text but preserve original
    clean_text = text.strip()
    if not clean_text:
        return text, "", ""
    
    # Check for text that's mostly punctuation and handle differently
    punctuation_count = sum(1 for c in clean_text if not c.isalnum() and not c.isspace())
    if punctuation_count > len(clean_text) / 2:
        # For text that's mostly punctuation, don't apply bionic formatting
        debug_print(f"  Skipping bionic formatting for punctuation-heavy text: '{text}'")
        return clean_text, "", clean_text
    
    # Determine how many characters to bold based on word length
    word_length = len(clean_text)
    if word_length <= 1:
        bold_length = 1
    elif word_length <= 3:
        bold_length = 1
    elif word_length <= 6:
        bold_length = 2
    elif word_length <= 9:
        bold_length = 3
    else:
        bold_length = min(4, word_length // 2)
    
    bold_part = clean_text[:bold_length]
    regular_part = clean_text[bold_length:]
    
    debug_print(f"  Bionic processing: '{clean_text}' → Bold: '{bold_part}', Regular: '{regular_part}'")
    return clean_text, bold_part, regular_part

def create_bionic_pdf_content_stream(input_path):
    """
    Create a bionic reading PDF by directly modifying the PDF content stream.
    This preserves the original document layout, images, and formatting while
    applying bionic reading enhancements.
    
    Args:
        input_path (str): Path to the input PDF file.
        
    Returns:
        str: Path to the output PDF file.
    """
    # Create output path
    base_name = os.path.splitext(os.path.basename(input_path))[0]
    temp_dir = tempfile.mkdtemp()
    output_path = os.path.join(temp_dir, f"{base_name}_bionic.pdf")
    
    try:
        # Open the document
        doc = fitz.open(input_path)
        debug_print(f"Processing PDF with {len(doc)} pages")
        
        # Process each page
        for page_num in range(len(doc)):
            page = doc[page_num]
            debug_print(f"Processing page {page_num + 1}")
            
            # Get all text blocks with their bounding boxes and properties
            text_blocks = page.get_text("dict")["blocks"]
            
            # Track our annotations for later highlighting
            bionic_annotations = []
            
            # Iterate through blocks, lines, and spans for fine-grained control
            for block_idx, block in enumerate(text_blocks):
                if "lines" not in block:
                    continue  # Skip non-text blocks
                
                for line_idx, line in enumerate(block["lines"]):
                    for span_idx, span in enumerate(line["spans"]):
                        # Extract span information
                        text = span["text"]
                        if not text.strip():
                            continue
                            
                        # Extract span properties for reconstruction
                        font_name = span["font"]
                        font_size = span["size"]
                        color = span["color"]
                        bbox = span["bbox"]
                        flags = span["flags"]
                        
                        debug_print(f"Processing text: '{text[:20]}...' with font={font_name}, size={font_size}")
                        
                        # Check for italic text - adjust our approach if italic
                        is_italic = (flags & 2) != 0
                        is_bold = (flags & 16) != 0
                        
                        debug_print(f"  Text style: italic={is_italic}, bold={is_bold}")
                        
                        # Parse the text into words while preserving spacing and punctuation
                        # This improved regex handles punctuation better and maintains word boundaries
                        words = []
                        for match in re.finditer(r'(\s+)|([^\s]+(?:\s*[^\w\s]\s*[^\s]+)*)', text):
                            if match.group(1):  # Space
                                words.append((match.group(1), None))
                            else:  # Word or word with punctuation
                                words.append((None, match.group(2)))
                        
                        debug_print(f"  Parsed into {len(words)} word/space segments: {[w for _, w in words if w][:5]}")
                        
                        # Create a replacement string with appropriate formatting
                        replaced_text = ""
                        replacement_positions = []
                        
                        current_pos = 0
                        for idx, (space, word) in enumerate(words):
                            # Handle spaces
                            if space:
                                replaced_text += space
                                current_pos += len(space)
                                continue
                            
                            # Skip empty words
                            if not word:
                                continue
                            
                            # Process word for bionic reading
                            _, bold_part, regular_part = process_text_bionic(word)
                            
                            if bold_part:
                                # Instead of using markup that might interfere with text rendering,
                                # we'll keep track of the positions to highlight separately
                                start_pos = current_pos
                                bold_length = len(bold_part)
                                end_pos = start_pos + bold_length
                                
                                # Store position of the bold part for later highlighting
                                replacement_positions.append((start_pos, end_pos))
                                
                                # Add the full word to the replaced text
                                replaced_text += bold_part + regular_part
                                current_pos += len(bold_part) + len(regular_part)
                                
                                debug_print(f"  Word: '{word}' → Bold: '{bold_part}', Regular: '{regular_part}'")
                            else:
                                # If processing failed, just add the original word
                                replaced_text += word
                                current_pos += len(word)
                                debug_print(f"  Word: '{word}' (unchanged)")
                        
                        # If this is italic or special text, use direct text extraction and insertion
                        # rather than the annotation approach which can cause issues with special fonts
                        if is_italic or len(replacement_positions) == 0:
                            debug_print(f"  Using direct text extraction for italic/special text")
                            
                            # Remove the original text with redaction
                            page.add_redact_annot(fitz.Rect(bbox), fill=(1, 1, 1))
                            page.apply_redactions()
                            
                            # Now insert the processed text with proper formatting
                            x_pos = bbox[0]
                            y_pos = bbox[1] + font_size * 0.8  # Adjust baseline
                            
                            # Split the text and position it manually for italic text
                            current_x = x_pos
                            current_pos = 0
                            
                            # Track last inserted character to help with spacing
                            last_char_type = None  # None, 'space', 'word', 'punct'
                            
                            debug_print(f"  Processing {len(words)} segments for direct insertion")
                            for idx, (space, word) in enumerate(words):
                                debug_print(f"  Segment {idx}: space={space is not None}, word={word}")
                                
                                if space:
                                    # Calculate space width - use a consistent approach
                                    space_width = len(space) * font_size * 0.5  # Increased width for better spacing
                                    current_x += space_width
                                    last_char_type = 'space'
                                    debug_print(f"  Added space width: {space_width}, current_x: {current_x}")
                                    continue
                                
                                if not word:
                                    continue
                                
                                # Check if this is a punctuation-only word
                                is_punctuation = all(not c.isalnum() and not c.isspace() for c in word)
                                
                                # Get the bold/regular parts
                                _, bold_part, regular_part = process_text_bionic(word)
                                
                                # If last inserted wasn't a space and this isn't punctuation,
                                # ensure there's spacing between words
                                if last_char_type == 'word' and not is_punctuation:
                                    debug_print(f"  Adding extra word spacing before: '{word}'")
                                    current_x += font_size * 0.2  # Add a small gap between words

                                if bold_part:
                                    try:
                                        # Insert bold part with bold formatting
                                        # For italic text, make sure to preserve the italic font
                                        # and use appropriate render mode
                                        render_mode = 0  # Default normal rendering
                                        
                                        if is_bold and is_italic:
                                            # Bold italic - ensure we're using a font that supports both
                                            if 'BoldOblique' in font_name or 'BoldItalic' in font_name:
                                                actual_font = font_name  # Font already supports both
                                            else:
                                                # Try to find a font that supports both
                                                actual_font = font_name.replace('Oblique', 'BoldOblique').replace('Italic', 'BoldItalic')
                                                if 'Bold' not in actual_font:
                                                    actual_font = actual_font.replace('Oblique', '').replace('Italic', '')
                                                    actual_font = actual_font + '-BoldOblique'
                                            render_mode = 2  # Fill + stroke helps with bold
                                        elif is_bold:
                                            # Just bold
                                            if 'Bold' in font_name:
                                                actual_font = font_name
                                            else:
                                                actual_font = font_name + '-Bold' if '-' not in font_name else font_name
                                            render_mode = 2  # Fill + stroke for bold effect
                                        elif is_italic:
                                            # Just italic
                                            if 'Oblique' in font_name or 'Italic' in font_name:
                                                actual_font = font_name
                                            else:
                                                actual_font = font_name + '-Oblique' if '-' not in font_name else font_name
                                            render_mode = 0  # Normal rendering for italic
                                        else:
                                            # Normal text - bold part should be bold
                                            actual_font = font_name
                                            render_mode = 2  # Make the first part bold with fill + stroke
                                        
                                        # Special case: if we're making normal text bold at beginning,
                                        # ensure we're using a bold font
                                        if not is_bold and render_mode == 2 and 'Bold' not in actual_font:
                                            actual_font = actual_font + '-Bold' if '-' not in actual_font else actual_font
                                        
                                        debug_print(f"  Bold part font: {actual_font}, render_mode: {render_mode}")
                                        
                                        page.insert_text(
                                            (current_x, y_pos),
                                            bold_part,
                                            fontname=actual_font,
                                            fontsize=font_size,
                                            fontfile=None,  # Let PyMuPDF handle font substitution
                                            color=color,
                                            fill=color,
                                            render_mode=render_mode
                                        )
                                        
                                        # Calculate width based on estimated width per character
                                        # Don't use text_length as it may not be available in this PyMuPDF version
                                        bold_width = len(bold_part) * font_size * 0.75
                                            
                                        current_x += bold_width
                                        
                                        # Track that we inserted a word character
                                        if is_punctuation:
                                            last_char_type = 'punct'
                                        else:
                                            last_char_type = 'word'
                                        
                                        # Add a small space after bold part if needed for readability
                                        # This helps with visual separation in italic text
                                        if is_italic and bold_part and regular_part:
                                            current_x += font_size * 0.15
                                        
                                    except Exception as e:
                                        debug_print(f"Error inserting bold text: {e}")
                                        # Fallback: insert the whole word without formatting
                                        try:
                                            page.insert_text(
                                                (current_x, y_pos),
                                                word,
                                                fontname=font_name,
                                                fontsize=font_size,
                                                color=color
                                            )
                                            
                                            word_width = len(word) * font_size * 0.75
                                            current_x += word_width
                                        except Exception as e2:
                                            debug_print(f"Fallback text insertion failed: {e2}")
                                        continue
                                
                                if regular_part:
                                    try:
                                        # Insert regular part with appropriate formatting
                                        # Make sure to preserve the original styling (italic, bold)
                                        render_mode = 0  # Default normal rendering
                                        
                                        # For the regular part, preserve the original styling
                                        if is_bold and is_italic:
                                            # Bold italic text should remain bold italic
                                            if 'BoldOblique' in font_name or 'BoldItalic' in font_name:
                                                actual_font = font_name
                                            else:
                                                actual_font = font_name.replace('Oblique', 'BoldOblique').replace('Italic', 'BoldItalic')
                                                if 'Bold' not in actual_font:
                                                    actual_font = actual_font.replace('Oblique', '').replace('Italic', '')
                                                    actual_font = actual_font + '-BoldOblique'
                                            render_mode = 0  # Normal rendering is enough as font has styling
                                        elif is_bold:
                                            # Bold text should remain bold
                                            if 'Bold' in font_name:
                                                actual_font = font_name
                                            else:
                                                actual_font = font_name + '-Bold' if '-' not in font_name else font_name
                                            render_mode = 0
                                        elif is_italic:
                                            # Italic text should remain italic
                                            if 'Oblique' in font_name or 'Italic' in font_name:
                                                actual_font = font_name
                                            else:
                                                actual_font = font_name + '-Oblique' if '-' not in font_name else font_name
                                            render_mode = 0
                                        else:
                                            # Normal text remains normal
                                            actual_font = font_name
                                            render_mode = 0
                                        
                                        debug_print(f"  Regular part font: {actual_font}, render_mode: {render_mode}")
                                        
                                        page.insert_text(
                                            (current_x, y_pos),
                                            regular_part,
                                            fontname=actual_font,
                                            fontsize=font_size,
                                            fontfile=None,
                                            color=color,
                                            fill=color,
                                            render_mode=render_mode
                                        )
                                        
                                        # Calculate width without using text_length
                                        regular_width = len(regular_part) * font_size * 0.75
                                        
                                        current_x += regular_width
                                        
                                        # Track that we inserted a word character
                                        if is_punctuation:
                                            last_char_type = 'punct'
                                        else:
                                            last_char_type = 'word'
                                        
                                    except Exception as e:
                                        debug_print(f"Error inserting regular text: {e}")
                                        debug_print(f"  Font info: {font_name}, is_italic={is_italic}, is_bold={is_bold}")
                                
                                # Add a small space between words if the next item isn't a space
                                if idx < len(words) - 1 and words[idx+1][0] is None:
                                    # If this is punctuation, add less space after it
                                    if is_punctuation:
                                        current_x += font_size * 0.15
                                    else:
                                        current_x += font_size * 0.3
                                # Always add a minimum spacing to prevent words from merging
                                elif not is_punctuation:  # Don't add extra space after punctuation
                                    current_x += font_size * 0.1
                        else:
                            # For normal text, use the redaction approach which works well
                            debug_print(f"  Using redaction approach for standard text")
                            
                            page.add_redact_annot(
                                fitz.Rect(bbox),
                                text=replaced_text,
                                fontname=font_name,
                                fontsize=font_size
                            )
                            
                            # Apply redactions
                            page.apply_redactions()
                            
                            # Now add highlighting for the bold parts
                            # We use a different approach based on text positions that's more reliable
                            for start_pos, end_pos in replacement_positions:
                                # Calculate the width of each part to highlight
                                highlight_text = replaced_text[start_pos:end_pos]
                                
                                # Find this text on the page with appropriate context
                                # Use an expanded rectangle for search instead of inflate which may not be available
                                rect = fitz.Rect(bbox)
                                # Expand the rectangle by 5 points in each direction
                                expanded_rect = fitz.Rect(rect.x0 - 5, rect.y0 - 5, rect.x1 + 5, rect.y1 + 5)
                                
                                try:
                                    instances = page.search_for(highlight_text, clip=expanded_rect)
                                    
                                    # Add highlight annotations for each instance
                                    for inst in instances:
                                        highlight = page.add_highlight_annot(inst)
                                        highlight.set_colors(stroke=(0.3, 0.3, 0.3))
                                        highlight.update()
                                        bionic_annotations.append(highlight)
                                except Exception as e:
                                    debug_print(f"Error highlighting text: {e}")
        
            debug_print(f"Page {page_num + 1} completed with {len(bionic_annotations)} highlight annotations")
        
        # Save the modified document
        doc.save(output_path)
        doc.close()
        
        debug_print(f"PDF processing complete, saved to {output_path}")
        return output_path
    
    except Exception as e:
        print(f"Error in content stream processing: {e}")
        import traceback
        traceback.print_exc()  # Print full error for debugging
        
        # If we encounter an error, fall back to more reliable methods
        try:
            return create_bionic_pdf_alternative(input_path)
        except Exception as e2:
            print(f"Alternative method also failed: {e2}")
            return create_bionic_pdf_reportlab(input_path)

def create_bionic_docx_direct(input_path):
    """
    Create a bionic reading DOCX by directly modifying the document content.
    
    Args:
        input_path (str): Path to the input DOCX file.
        
    Returns:
        str: Path to the output DOCX file.
    """
    # Create output path
    base_name = os.path.splitext(os.path.basename(input_path))[0]
    temp_dir = tempfile.mkdtemp()
    output_path = os.path.join(temp_dir, f"{base_name}_bionic.docx")
    
    # Open the document
    doc = Document(input_path)
    
    # Better approach: Process runs within paragraphs
    for paragraph in doc.paragraphs:
        # Create a list to store new runs
        new_runs_data = []
        
        # Process each run
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
                
            # Store run formatting
            font = run.font
            
            # Process words in the run
            words = re.findall(r'\S+|\s+', text)
            
            for word in words:
                if word.isspace():
                    # Preserve whitespace
                    new_runs_data.append({
                        'text': word,
                        'bold': font.bold,
                        'italic': font.italic,
                        'size': font.size,
                        'color': font.color.rgb if font.color and font.color.rgb else None,
                        'name': font.name
                    })
                else:
                    # Process word for bionic reading
                    clean_word, bold_part, regular_part = process_text_bionic(word)
                    
                    if bold_part:
                        # Add bold part as separate run
                        new_runs_data.append({
                            'text': bold_part,
                            'bold': True,  # Make it bold
                            'italic': font.italic,
                            'size': font.size,
                            'color': font.color.rgb if font.color and font.color.rgb else None,
                            'name': font.name
                        })
                    
                    if regular_part:
                        # Add regular part
                        new_runs_data.append({
                            'text': regular_part,
                            'bold': font.bold,  # Keep original bold setting
                            'italic': font.italic,
                            'size': font.size,
                            'color': font.color.rgb if font.color and font.color.rgb else None,
                            'name': font.name
                        })
        
        # Clear existing runs
        paragraph.clear()
        
        # Add new runs with bionic formatting
        for run_data in new_runs_data:
            new_run = paragraph.add_run(run_data['text'])
            
            # Apply formatting
            if run_data['bold'] is not None:
                new_run.font.bold = run_data['bold']
            if run_data['italic'] is not None:
                new_run.font.italic = run_data['italic']
            if run_data['size'] is not None:
                new_run.font.size = run_data['size']
            if run_data['color'] is not None:
                new_run.font.color.rgb = run_data['color']
            if run_data['name'] is not None:
                new_run.font.name = run_data['name']
    
    # Save the modified document
    doc.save(output_path)
    
    return output_path

def create_bionic_text_file(input_path):
    """
    Create a bionic reading text file using HTML formatting.
    
    Args:
        input_path (str): Path to the input text file.
        
    Returns:
        str: Path to the output HTML file.
    """
    # Read the text file
    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create output path
    base_name = os.path.splitext(os.path.basename(input_path))[0]
    temp_dir = tempfile.mkdtemp()
    output_path = os.path.join(temp_dir, f"{base_name}_bionic.html")
    
    # Process text for bionic reading
    def process_text_html(text):
        words = re.findall(r'\b\w+\b|\W+', text)
        result = []
        
        for word in words:
            if not re.match(r'\b\w+\b', word):
                result.append(word)
                continue
                
            clean_word, bold_part, regular_part = process_text_bionic(word)
            
            if bold_part and regular_part:
                result.append(f'<strong>{bold_part}</strong>{regular_part}')
            elif bold_part:
                result.append(f'<strong>{bold_part}</strong>')
            else:
                result.append(word)
        
        return ''.join(result)
    
    # Create HTML content with improved styling
    html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Bionic Reading - {os.path.basename(input_path)}</title>
    <style>
        body {{
            font-family: 'Georgia', 'Times New Roman', serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            background-color: #fafafa;
            color: #333;
        }}
        .container {{
            background-color: white;
            padding: 40px;
            border-radius: 8px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }}
        h1 {{
            color: #2c3e50;
            border-bottom: 2px solid #3498db;
            padding-bottom: 10px;
        }}
        p {{
            margin-bottom: 1em;
            text-align: justify;
        }}
        strong {{
            font-weight: bold;
            color: #2c3e50;
        }}
        pre {{
            white-space: pre-wrap;
            font-family: inherit;
            margin: 0;
        }}
        .footer {{
            margin-top: 30px;
            text-align: center;
            font-size: 0.8rem;
            color: #777;
        }}
    </style>
</head>
<body>
    <div class="container">
        <h1>Bionic Reading Format</h1>
        <h2>Original file: {os.path.basename(input_path)}</h2>
"""
    
    # Improved paragraph detection:
    # 1. Split by double newline (traditional paragraphs)
    # 2. Also handle single newline content with proper line breaks
    
    # First try to detect if this is a paragraph-based text or line-by-line text
    if content.count('\n\n') > 0:
        # This is likely a paragraph-based document
        paragraphs = content.split('\n\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                # Handle any single line breaks within paragraphs
                lines = paragraph.strip().split('\n')
                if len(lines) > 1:
                    # Paragraph has internal line breaks
                    processed_lines = [process_text_html(line) for line in lines if line.strip()]
                    html_content += f"        <p>{processed_lines[0]}"
                    for line in processed_lines[1:]:
                        html_content += f"<br>\n        {line}"
                    html_content += "</p>\n"
                else:
                    # Simple paragraph
                    processed_paragraph = process_text_html(paragraph.strip())
                    html_content += f"        <p>{processed_paragraph}</p>\n"
    else:
        # This is likely a line-by-line document
        lines = content.split('\n')
        if len(lines) > 1:
            html_content += "        <pre>\n"
            for line in lines:
                if line.strip():
                    processed_line = process_text_html(line)
                    html_content += f"        {processed_line}\n"
                else:
                    html_content += "\n"
            html_content += "        </pre>\n"
        else:
            # Single line document
            processed_content = process_text_html(content.strip())
            html_content += f"        <p>{processed_content}</p>\n"
    
    html_content += """        <div class="footer">
            <p>Processed with ZapRead Bionic Reading</p>
        </div>
    </div>
</body>
</html>"""
    
    # Save HTML file
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html_content)
    
    return output_path

def detect_file_type(file_path):
    """
    Detect the file type of a given file.
    
    Args:
        file_path (str): Path to the file.
        
    Returns:
        str: The detected file type (txt, pdf, docx, or unknown).
    """
    try:
        mime = magic.Magic(mime=True)
        file_type = mime.from_file(file_path)
        
        if file_type == 'text/plain':
            return 'txt'
        elif file_type == 'application/pdf':
            return 'pdf'
        elif file_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                           'application/msword']:
            return 'docx'
        else:
            # Fallback to extension-based detection
            ext = os.path.splitext(file_path)[1].lower()
            if ext == '.txt':
                return 'txt'
            elif ext == '.pdf':
                return 'pdf'
            elif ext in ['.docx', '.doc']:
                return 'docx'
            else:
                return 'unknown'
    except Exception:
        # Fallback to extension-based detection
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.txt':
            return 'txt'
        elif ext == '.pdf':
            return 'pdf'
        elif ext in ['.docx', '.doc']:
            return 'docx'
        else:
            return 'unknown'

def create_bionic_pdf_reportlab(input_path):
    """
    Create a bionic reading PDF using ReportLab for maximum compatibility.
    This is a fallback method when the primary method fails.
    """
    base_name = os.path.splitext(os.path.basename(input_path))[0]
    temp_dir = tempfile.mkdtemp()
    output_path = os.path.join(temp_dir, f"{base_name}_bionic.pdf")
    
    # Try to extract text from PDF using pdfplumber
    extracted_text = []
    try:
        with pdfplumber.open(input_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    extracted_text.append(text)
    except Exception as e:
        print(f"Error extracting text with pdfplumber: {e}")
        # Fall back to PyMuPDF for text extraction
        try:
            doc = fitz.open(input_path)
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                if text:
                    extracted_text.append(text)
            doc.close()
        except Exception as e2:
            print(f"Error extracting text with PyMuPDF: {e2}")
            # If all extraction methods fail, return empty list
            extracted_text = ["Failed to extract text from PDF"]
    
    # Create PDF using ReportLab
    buffer = io.BytesIO()
    
    try:
        # Use built-in PDF fonts instead of trying to register external TTF fonts
        # Create the PDF document
        c = canvas.Canvas(buffer, pagesize=letter)
        width, height = letter
        
        # Set up text formatting
        font_size = 12
        line_height = font_size * 1.2
        margin = 72  # 1 inch margin
        
        # Use standard PDF fonts that are guaranteed to be available
        normal_font = "Helvetica"
        bold_font = "Helvetica-Bold"
        
        for page_text in extracted_text:
            y = height - margin
            
            # Process paragraphs
            paragraphs = page_text.split('\n\n')
            for paragraph in paragraphs:
                if not paragraph.strip():
                    continue
                
                # Process lines
                lines = paragraph.split('\n')
                for line in lines:
                    if not line.strip():
                        y -= line_height
                        continue
                    
                    # Process words
                    words = line.split()
                    x = margin
                    
                    for word in words:
                        # Process word for bionic reading
                        clean_word, bold_part, regular_part = process_text_bionic(word)
                        
                        if bold_part:
                            # Set font for bold part
                            c.setFont(bold_font, font_size)
                            
                            # Draw bold part
                            c.drawString(x, y, bold_part)
                            
                            # Calculate width of bold part
                            bold_width = c.stringWidth(bold_part, bold_font, font_size)
                            x += bold_width
                        
                        if regular_part:
                            # Set font for regular part
                            c.setFont(normal_font, font_size)
                            
                            # Draw regular part
                            c.drawString(x, y, regular_part)
                            
                            # Calculate width of regular part
                            regular_width = c.stringWidth(regular_part, normal_font, font_size)
                            x += regular_width
                        
                        # Add space between words
                        x += c.stringWidth(" ", normal_font, font_size)
                    
                    y -= line_height
                    
                    # Check if we need a new page
                    if y < margin:
                        c.showPage()
                        y = height - margin
                
                # Add space between paragraphs
                y -= line_height
            
            # End the page
            c.showPage()
        
        # Save the PDF
        c.save()
        
        # Get the PDF content and write to file
        buffer.seek(0)
        with open(output_path, 'wb') as f:
            f.write(buffer.getvalue())
        
        return output_path
    
    except Exception as e:
        print(f"ReportLab PDF generation failed: {e}")
        # Last resort - try the basic method
        return create_bionic_pdf_basic(input_path)

# Basic fallback method
def create_bionic_pdf_basic(input_path):
    """
    Very basic PDF creation method that uses minimal dependencies and features.
    This is the last resort fallback.
    """
    base_name = os.path.splitext(os.path.basename(input_path))[0]
    temp_dir = tempfile.mkdtemp()
    output_path = os.path.join(temp_dir, f"{base_name}_bionic.pdf")
    
    try:
        # Extract text using PyMuPDF
        doc = fitz.open(input_path)
        text_content = ""
        
        for page_num in range(len(doc)):
            text_content += doc[page_num].get_text()
        
        doc.close()
        
        # Create a very basic PDF with minimal formatting
        pdf = canvas.Canvas(output_path, pagesize=letter)
        width, height = letter
        
        # Process text for bionic reading
        processed_text = ""
        lines = text_content.split('\n')
        
        for line in lines:
            if not line.strip():
                processed_text += '\n'
                continue
                        
            words = line.split()
            line_parts = []
            
            for word in words:
                clean_word, bold_part, regular_part = process_text_bionic(word)
                if bold_part and regular_part:
                    # We can't do bold in this simple format, so we just combine them
                    line_parts.append(bold_part + regular_part)
                else:
                    line_parts.append(word)
                    
            processed_text += ' '.join(line_parts) + '\n'
        
        # Write text to PDF
        y = height - 72  # Start 1 inch from top
        for line in processed_text.split('\n'):
            if y < 72:  # If less than 1 inch from bottom, start new page
                pdf.showPage()
                y = height - 72
                
            pdf.drawString(72, y, line)
            y -= 15  # Move down by 15 points
            
        pdf.save()
        
        return output_path
    
    except Exception as e:
        print(f"Basic PDF creation failed: {e}")
        # As a very last resort, create an HTML file instead
        html_output = os.path.join(temp_dir, f"{base_name}_bionic.html")
        
        with open(html_output, 'w', encoding='utf-8') as f:
            f.write(f"""<!DOCTYPE html>
<html>
<head>
    <title>Bionic Reading - {base_name}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
        .error {{ color: red; font-weight: bold; }}
        .content {{ background-color: #f9f9f9; padding: 20px; border-radius: 5px; }}
    </style>
</head>
<body>
    <h1>Bionic Reading</h1>
    <div class="error">
        <p>The PDF could not be processed in bionic format due to technical issues.</p>
        <p>Please try another file format or contact support.</p>
    </div>
    <div class="content">
        <p>Original filename: {os.path.basename(input_path)}</p>
    </div>
</body>
</html>""")
        
        return html_output

def create_bionic_pdf_alternative(input_path):
    """
    An alternative implementation that creates a bionic reading PDF.
    This version aims to preserve page dimensions, font sizes, basic text styling, and images.
    """
    base_name = os.path.splitext(os.path.basename(input_path))[0]
    temp_dir = tempfile.mkdtemp()
    output_path = os.path.join(temp_dir, f"{base_name}_bionic.pdf")

    try:
        doc = fitz.open(input_path)
        all_page_data = []

        # 1. Extract data from each page
        for page_num in range(len(doc)):
            src_page = doc[page_num]
            page_width = src_page.rect.width
            page_height = src_page.rect.height

            # Extract images
            page_images_data = []
            img_list = src_page.get_images(full=True)
            for img_index, img_info in enumerate(img_list):
                xref = img_info[0]
                base_image = doc.extract_image(xref)
                if not base_image:
                    continue
                
                img_bytes = base_image["image"]
                img_ext = base_image["ext"]
                
                img_rects = src_page.get_image_rects(img_info) 
                if img_rects:
                    bbox_pdf = img_rects[0] 
                    page_images_data.append({
                        "bytes": img_bytes,
                        "ext": img_ext,
                        "bbox_pdf": fitz.Rect(bbox_pdf.x0, bbox_pdf.y0, bbox_pdf.x1, bbox_pdf.y1)
                    })

            # Extract text spans
            page_text_spans = []
            text_dict = src_page.get_text("dict") # Use default flags
            
            for block in text_dict["blocks"]:
                if block["type"] == 0: # Text block
                    for line in block["lines"]:
                        for span in line["spans"]:
                            if span["text"].strip():
                                page_text_spans.append({
                                    "text": span["text"],
                                    "bbox_pdf": fitz.Rect(span["bbox"]),
                                    "origin_pdf": fitz.Point(span["origin"]),
                                    "font_size": span["size"],
                                    "flags": span["flags"],
                                    "font_name": span["font"]
                                })

            all_page_data.append({
                "width": page_width,
                "height": page_height,
                "images": page_images_data,
                "text_spans": page_text_spans,
            })

        doc.close()

        # 2. Generate PDF with ReportLab
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer) 

        for page_data in all_page_data:
            page_width_rl = page_data["width"]
            page_height_rl = page_data["height"]
            c.setPageSize((page_width_rl, page_height_rl))

            # Draw Images
            for img_info in page_data["images"]:
                img_bytes = img_info["bytes"]
                bbox_pdf = img_info["bbox_pdf"]

                rl_img_x = bbox_pdf.x0
                rl_img_y = page_height_rl - bbox_pdf.y1 
                rl_img_width = bbox_pdf.width
                rl_img_height = bbox_pdf.height
                
                try:
                    image_reader = ImageReader(io.BytesIO(img_bytes))
                    c.drawImage(image_reader, rl_img_x, rl_img_y, 
                                width=rl_img_width, height=rl_img_height, mask='auto')
                except Exception as e:
                    print(f"Error drawing image: {e}. Image bbox: {bbox_pdf}")

            # Draw Text
            for span_info in page_data["text_spans"]:
                text = span_info["text"]
                font_size = span_info["font_size"]
                flags = span_info["flags"]
                # bbox_pdf = span_info["bbox_pdf"] # Not directly used for x,y if origin is used
                origin_pdf = span_info["origin_pdf"]

                rl_x_start = origin_pdf.x
                rl_y_baseline = page_height_rl - origin_pdf.y
                
                current_x = rl_x_start

                words = re.findall(r'\S+|\s+', text)
                
                for i, word in enumerate(words):
                    if not word.strip(): 
                        space_width = c.stringWidth(word, "Helvetica", font_size) if word else 0
                        current_x += space_width
                        continue

                    clean_word, bold_part, regular_part = process_text_bionic(word)

                    if bold_part:
                        c.setFont("Helvetica-Bold", font_size)
                        c.drawString(current_x, rl_y_baseline, bold_part)
                        current_x += c.stringWidth(bold_part, "Helvetica-Bold", font_size)
                    
                    if regular_part:
                        is_bold_flag = flags & 16
                        is_italic_flag = flags & 2
                        
                        font_rl = "Helvetica"
                        if is_bold_flag and is_italic_flag:
                            font_rl = "Helvetica-BoldOblique"
                        elif is_bold_flag:
                            font_rl = "Helvetica-Bold"
                        elif is_italic_flag:
                            font_rl = "Helvetica-Oblique"
                        
                        c.setFont(font_rl, font_size)
                        c.drawString(current_x, rl_y_baseline, regular_part)
                        current_x += c.stringWidth(regular_part, font_rl, font_size)
                    
                    # Add a small space after processing a non-space word, before the next word (if any)
                    if i < len(words) - 1: # Ensure it's not the last word token
                        # Check if the original text likely had a space here by looking at the next token
                        if words[i+1].isspace():
                            # If next token is space, it will be handled by the space logic, so do nothing here
                            pass 
                        else: # If next token is another word, add a standard small space
                            current_x += c.stringWidth(" ", "Helvetica", font_size) * 0.3 
            
            c.showPage()

        c.save()
        
        with open(output_path, 'wb') as f:
            f.write(buffer.getvalue())
    
        return output_path

    except Exception as e:
        print(f"Create_bionic_pdf_alternative processing failed: {e}")
        import traceback
        traceback.print_exc() # Print full traceback for better debugging
        try:
            print(f"Falling back to create_bionic_pdf_reportlab due to error: {e}")
            return create_bionic_pdf_reportlab(input_path)
        except Exception as e_fallback:
            print(f"Fallback create_bionic_pdf_reportlab also failed: {e_fallback}")
            return create_bionic_pdf_basic(input_path) # Final fallback

def process_file_with_fallback(file_path):
    """
    Process file with multiple fallback options for better reliability.
    """
    file_type = detect_file_type(file_path)
    debug_print(f"Processing file: {file_path}")
    debug_print(f"Detected file type: {file_type}")
    
    if file_type == 'pdf':
        try:
            print("Starting PDF processing with content stream method...")
            debug_print("Attempting content stream method first...")
            return create_bionic_pdf_content_stream(file_path), file_type, True
        except Exception as e:
            print(f"Content stream PDF processing failed: {e}, trying alternative method...")
            debug_print(f"Content stream method failed with error: {e}")
            import traceback
            debug_print(traceback.format_exc())
            
            try:
                print("Trying PDF processing with Alternative Structure method...")
                debug_print("Attempting alternative structure method...")
                return create_bionic_pdf_alternative(file_path), file_type, True
            except Exception as e2:
                print(f"Alternative PDF processing failed: {e2}, trying ReportLab method...")
                debug_print(f"Alternative method failed with error: {e2}")
                debug_print(traceback.format_exc())
                
                try:
                    debug_print("Attempting ReportLab method...")
                    return create_bionic_pdf_reportlab(file_path), file_type, True
                except Exception as e3:
                    print(f"ReportLab PDF processing failed: {e3}, trying basic method...")
                    debug_print(f"ReportLab method failed with error: {e3}")
                    debug_print(traceback.format_exc())
                    
                    try:
                        debug_print("Attempting basic method as last resort...")
                        return create_bionic_pdf_basic(file_path), file_type, True
                    except Exception as e4:
                        print(f"All PDF processing methods failed: {e4}")
                        debug_print(f"All methods failed with final error: {e4}")
                        debug_print(traceback.format_exc())
                        return None, file_type, False
    
    elif file_type == 'docx':
        try:
            debug_print("Processing DOCX file...")
            return create_bionic_docx_direct(file_path), file_type, True
        except Exception as e:
            print(f"DOCX processing failed: {e}")
            debug_print(f"DOCX processing failed with error: {e}")
            import traceback
            debug_print(traceback.format_exc())
            return None, file_type, False
    
    elif file_type == 'txt':
        try:
            debug_print("Processing TXT file...")
            return create_bionic_text_file(file_path), file_type, True
        except Exception as e:
            print(f"Text processing failed: {e}")
            debug_print(f"TXT processing failed with error: {e}")
            import traceback
            debug_print(traceback.format_exc())
            return None, file_type, False
    
    else:
        debug_print(f"Unsupported file type detected: {file_type}")
        return None, file_type, False

# ===========================================================================
# ADAPTER FUNCTIONS TO MAINTAIN COMPATIBILITY WITH EXISTING APP CODE
# ===========================================================================

def process_file(file_path):
    """
    Adapter function for compatibility with the existing app interface.
    
    Args:
        file_path (str): Path to the file.
        
    Returns:
        tuple: (processed_content, file_type, temp_dir)
    """
    # Use the enhanced version with fallbacks
    output_path, file_type, success = process_file_with_fallback(file_path)
    
    if not success or not output_path:
        raise ValueError(f"Failed to process file: {file_path}")
    
    # For binary files, return empty content but pass the path in temp_dir
    temp_dir = os.path.dirname(output_path)
    
    # Read the content of the processed file (for compatibility)
    with open(output_path, 'rb') as f:
        processed_content = f.read()
    
    # Return in the format expected by app.py
    return processed_content, file_type, temp_dir

def save_processed_file(content, original_filename, file_type, temp_dir=None):
    """
    Adapter function for compatibility with the existing app interface.
    This function is a simplified wrapper since the actual processing is now
    done in the process_file_with_fallback function.
    
    Args:
        content (bytes/str): The processed content from process_file.
        original_filename (str): The original filename.
        file_type (str): The original file type.
        temp_dir (str, optional): Temporary directory containing the processed file.
        
    Returns:
        str: Path to the processed file.
    """
    # With the new approach, the file is already saved during processing
    # We just need to find the file in the temp_dir and return its path
    if temp_dir and os.path.exists(temp_dir):
        base_name = os.path.splitext(os.path.basename(original_filename))[0]
        
        # Look for the generated file in the temp directory
        if file_type == 'pdf':
            expected_path = os.path.join(temp_dir, f"{base_name}_bionic.pdf")
            if os.path.exists(expected_path):
                return expected_path
        elif file_type == 'docx':
            expected_path = os.path.join(temp_dir, f"{base_name}_bionic.docx")
            if os.path.exists(expected_path):
                return expected_path
        elif file_type == 'txt':
            expected_path = os.path.join(temp_dir, f"{base_name}_bionic.html")
            if os.path.exists(expected_path):
                return expected_path
    
    # If the file wasn't found, or temp_dir is None, we need to create it
    # This fallback shouldn't normally be needed with the new implementation
    output_temp_dir = temp_dir if temp_dir else tempfile.mkdtemp()
    
    if file_type in ['pdf', 'docx']:
        extension = file_type
    else:
        extension = 'html'  # HTML output for text files
        
    base_name = os.path.splitext(os.path.basename(original_filename))[0]
    output_filename = f"{base_name}_bionic.{extension}"
    output_path = os.path.join(output_temp_dir, output_filename)
    
    # Save the content to the file
    with open(output_path, 'wb') as f:
        if isinstance(content, str):
            f.write(content.encode('utf-8'))
        else:
            f.write(content)
    
    return output_path